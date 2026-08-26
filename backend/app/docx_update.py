"""Write corrected text back into an ALREADY FORMATTED .docx.

Exporting a fresh document is the easy path and the wrong one: the team's
script carries a cover sheet (ใบปะหน้า), superscript story markers, tab
indents and hand-placed double spaces, and a fresh export throws every bit of
it away. What they asked for is the opposite — keep their document, change
only the words that actually changed (2026-08-25).

So nothing here rewrites a paragraph wholesale. Runs are edited in place at
the character level, which is what preserves formatting: a run nobody touched
keeps its own font, its superscript flag, its spacing, because it is never
written to at all.

Two rules earn their keep, both learned the hard way on this file:

- Content is found by MATCHING TEXT, never by counting paragraphs or looking
  for a marker. "เนื้อหาไม่ได้เริ่มที่ตัวคั่นตัวแรกนะคะ" — the script does not
  begin at the first marker, and any structural rule of that shape is wrong on
  the real document.
- Anything not confidently paired is LEFT ALONE. The cover sheet matches no
  script line, and a rule that deleted unmatched paragraphs would delete it.
  Deletion is allowed only strictly between two matched paragraphs, where the
  surrounding anchors prove the gap really is missing content.
"""

from __future__ import annotations

import copy
import os
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

# Comparing whole 250-character paragraphs to find a pairing is both slow and
# needlessly picky; a prefix is enough to tell "this is the same paragraph,
# edited" from "this is a different paragraph".
_HEAD_CHARS = 120
# Below this the two paragraphs are different content, not an edit of one
# another. Deliberately generous: a heavily rewritten line is still that line,
# and the cost of pairing too eagerly (an edit applied to the right paragraph)
# is far smaller than the cost of pairing too little (a paragraph left stale).
_PAIR_RATIO = 0.5
# Two paragraphs whose lengths differ by more than this are not an edit of one
# another, whatever their characters happen to share.
_LENGTH_RATIO = 0.4
# How much of the script the document must actually account for before this is
# treated as the same programme at all. One paragraph in common is not enough:
# programmes share stock openings and closings, so picking the wrong episode
# gave a single exact anchor, a "match", and then the whole of a different
# episode inserted after it (raised in review 2026-08-26). Measured on the
# team's own file a genuine pairing covers ~100% of the script, so half is
# generous — it still admits a document whose second half was rewritten.
_MIN_COVERAGE = 0.5


def similarity(a: str, b: str) -> float:
    """How alike two paragraphs are, on their opening characters.

    Length is checked first, and it is not a nicety. Thai shares so many
    characters between unrelated sentences that a short line can score over
    the pairing threshold against a long paragraph it has nothing to do with —
    which is how a completely unrelated script came close to being written
    over the team's document.
    """
    if not a or not b:
        return 0.0
    if min(len(a), len(b)) < _LENGTH_RATIO * max(len(a), len(b)):
        return 0.0
    return SequenceMatcher(None, a[:_HEAD_CHARS], b[:_HEAD_CHARS]).ratio()


@dataclass
class Change:
    """One difference, in the vocabulary the team already uses."""

    kind: str  # "แก้ไข" | "เพิ่มเติม" | "ตัดออก"
    paragraph: int  # index in the document (for ตัดออก/แก้ไข), or the anchor
    before: str
    after: str

    def to_dict(self) -> dict:
        return {
            "kind": self.kind,
            "paragraph": self.paragraph,
            "before": self.before,
            "after": self.after,
        }


def _run_spans(texts: list[str]) -> list[tuple[int, int]]:
    spans, pos = [], 0
    for t in texts:
        spans.append((pos, pos + len(t)))
        pos += len(t)
    return spans


def plan_runs(
    run_texts: list[str], new_text: str, protected: "set[int] | tuple[int, ...]" = ()
) -> list[tuple[int, str]]:
    """Lay out `new_text` as (which run's formatting, text) segments.

    Pure, so every placement rule below can be tested without a Word document.

    Unchanged characters stay with the run they were already in — that is the
    whole point, since a run that is never rewritten keeps its font, size and
    spacing untouched.

    `protected` names runs that must never absorb text that is not theirs.
    These are the superscript story markers. The first attempt simply appended
    inserted text to the run before it, and on this team's file that was wrong
    in the most damaging way possible: their markers sit at the END of a
    paragraph as often as the start, so words added to the end of a sentence
    landed INSIDE the marker and came out raised. Text is instead handed to
    the nearest ordinary run, and because a segment can name a run that is
    already in use, the marker keeps its place in the middle and the new words
    get a run of their own after it — right characters, right order, right
    formatting.
    """
    if not run_texts:
        return []
    old = "".join(run_texts)
    spans = _run_spans(run_texts)
    prot = set(protected)

    def run_at(pos: int) -> int:
        for k, (a, b) in enumerate(spans):
            if a <= pos < b:
                return k
        return len(run_texts) - 1

    def home(k: int) -> int:
        """The nearest run allowed to hold text that is not already there."""
        if k not in prot:
            return k
        for back in range(k - 1, -1, -1):
            if back not in prot:
                return back
        for fwd in range(k + 1, len(run_texts)):
            if fwd not in prot:
                return fwd
        return k

    segs: list[tuple[int, str]] = []

    def emit(src: int, text: str) -> None:
        if not text:
            return
        if segs and segs[-1][0] == src:
            segs[-1] = (src, segs[-1][1] + text)
        else:
            segs.append((src, text))

    for tag, i1, i2, j1, j2 in SequenceMatcher(
        None, old, new_text, autojunk=False
    ).get_opcodes():
        if tag == "equal":
            for off in range(i2 - i1):
                emit(run_at(i1 + off), new_text[j1 + off])
        elif tag == "replace":
            emit(home(run_at(i1)), new_text[j1:j2])
        elif tag == "insert":
            emit(home(run_at(i1 - 1) if i1 > 0 else 0), new_text[j1:j2])
        # "delete": those characters simply go nowhere
    return segs


def _protected_runs(runs: list) -> list[int]:
    """Runs whose formatting carries meaning of its own, not just style.

    Superscript and subscript are the team's story markers (¹[ … ]¹). Losing
    one, or raising a word that should sit on the line, is a visible defect in
    a document that goes out to the client.
    """
    return [k for k, r in enumerate(runs) if r.font.superscript or r.font.subscript]


def set_paragraph_text(par, new_text: str) -> bool:
    """Edit a paragraph's runs in place. Returns whether anything changed."""
    runs = list(par.runs)
    if not runs:
        return False
    old_texts = [r.text for r in runs]
    if "".join(old_texts) == new_text:
        return False

    plan = plan_runs(old_texts, new_text, _protected_runs(runs))
    if not plan:  # every character deleted — empty the runs, keep the shape
        for run in runs:
            run.text = ""
        return True

    # Same number of runs in the same order: only the text differs, so no XML
    # surgery is needed and nothing can be disturbed.
    if [src for src, _ in plan] == list(range(len(runs))):
        for run, (_, text) in zip(runs, plan):
            if run.text != text:
                run.text = text
        return True

    if any(run._r.getparent() is not par._p for run in runs):
        # A run nested inside a hyperlink or field: reordering those elements
        # is not safe, so fall back to placing text in the runs as they stand.
        for run, text in zip(runs, _flatten(plan, len(runs))):
            if run.text != text:
                run.text = text
        return True

    _rebuild_runs(par, runs, plan)
    return True


def _flatten(plan: list[tuple[int, str]], count: int) -> list[str]:
    out = [""] * count
    for src, text in plan:
        out[src] += text
    return out


def _rebuild_runs(par, runs: list, plan: list[tuple[int, str]]) -> None:
    """Rewrite the paragraph as the planned run sequence.

    A segment may name a run that has already been placed — that is how added
    text gets its own run with an existing run's formatting. The first use
    moves the original element; later uses copy it, so the formatting is
    carried without inventing any.
    """
    from docx.text.run import Run

    used: set[int] = set()
    planned = []
    for src, text in plan:
        if src in used:
            element = copy.deepcopy(runs[src]._r)
        else:
            element = runs[src]._r
            used.add(src)
        planned.append((element, text))

    for run in runs:
        par._p.remove(run._r)
    for element, text in planned:
        par._p.append(element)
        Run(element, par).text = text


def pair_up(doc_texts: list[str], lines: list[str]) -> list[tuple[int | None, int | None]]:
    """Pair document paragraphs with script lines, keeping their order.

    Identical paragraphs anchor the pass — most of them are identical, since
    only some lines were corrected — and the runs between anchors are paired
    by similarity. A pairing is never allowed to cross an anchor, so a local
    mess cannot throw the rest of the document out of step.

    Returns (paragraph index | None, line index | None) pairs: None on the
    left means a line with no paragraph (เพิ่มเติม), None on the right means a
    paragraph with no line (ตัดออก, or simply not part of the script).
    """
    pairs: list[tuple[int | None, int | None]] = []
    opcodes = SequenceMatcher(None, doc_texts, lines, autojunk=False).get_opcodes()
    anchors = [n for n, op in enumerate(opcodes) if op[0] == "equal"]
    first_anchor = anchors[0] if anchors else None
    last_anchor = anchors[-1] if anchors else None

    for n, (tag, i1, i2, j1, j2) in enumerate(opcodes):
        # Is this block hemmed in by paragraphs that matched exactly? Only
        # then is "these must be each other, rewritten" a safe conclusion.
        bracketed = (
            first_anchor is not None and first_anchor < n < last_anchor
        )
        if tag == "equal":
            pairs.extend((i, j) for i, j in zip(range(i1, i2), range(j1, j2)))
            continue
        if tag == "delete":
            pairs.extend((i, None) for i in range(i1, i2))
            continue
        if tag == "insert":
            pairs.extend((None, j) for j in range(j1, j2))
            continue
        # "replace": same stretch of the script, rewritten. Walk both sides
        # together and pair while they still look like each other.
        i, j = i1, j1
        while i < i2 and j < j2:
            if similarity(doc_texts[i], lines[j]) >= _PAIR_RATIO:
                pairs.append((i, j))
                i, j = i + 1, j + 1
                continue
            # Not a match. Whichever side finds its partner sooner is the one
            # that has extra entries here.
            skip_doc = next(
                (k for k in range(i + 1, i2) if similarity(doc_texts[k], lines[j]) >= _PAIR_RATIO),
                None,
            )
            skip_line = next(
                (k for k in range(j + 1, j2) if similarity(doc_texts[i], lines[k]) >= _PAIR_RATIO),
                None,
            )
            if skip_doc is not None and (skip_line is None or skip_doc - i <= skip_line - j):
                pairs.extend((k, None) for k in range(i, skip_doc))
                i = skip_doc
            elif skip_line is not None:
                pairs.extend((None, k) for k in range(j, skip_line))
                j = skip_line
            elif bracketed:
                # Neither side recognises the other, but identical paragraphs
                # sit on both sides of this block, so it really is this stretch
                # of the script, rewritten past recognition.
                pairs.append((i, j))
                i, j = i + 1, j + 1
            else:
                # No anchors around it. Pairing here is a guess, and the cost
                # of guessing wrong is the team's document overwritten with
                # text from a different programme — so refuse, and let the
                # caller see that nothing matched.
                break
        pairs.extend((k, None) for k in range(i, i2))
        pairs.extend((None, k) for k in range(j, j2))
    return pairs


def coverage(
    pairs: list[tuple[int | None, int | None]],
    doc_texts: list[str],
    lines: list[str],
) -> float:
    """Fraction of the script's characters that found a home in the document.

    Counting CHARACTERS, not paragraphs, is the point: one short shared line
    out of a hundred is a coincidence, and paragraph counting cannot tell that
    from a real match the way length can.
    """
    total = sum(len(line) for line in lines)
    if total == 0:
        return 0.0
    matched = sum(len(lines[j]) for i, j in pairs if i is not None and j is not None)
    return matched / total


def content_bounds(pairs: list[tuple[int | None, int | None]]) -> tuple[int, int] | None:
    """The paragraph range the script actually occupies.

    Everything outside it — the cover sheet above, anything below — is not
    ours to touch. This is why the cover sheet survives: it matches no script
    line, so it never enters the range, so no rule can reach it.
    """
    matched = [i for i, j in pairs if i is not None and j is not None]
    if not matched:
        return None
    return min(matched), max(matched)


def save_atomically(doc, out_path: Path) -> None:
    """Write the document without ever leaving the destination half-written.

    out_path is very often the team's own master file — the dialog offers
    "save over the original" on purpose — and python-docx writes a ZIP in
    place. A crash, a full disk or an antivirus lock partway through would
    leave them with a .docx Word cannot open and no copy to fall back on.
    The audio export already writes to .part and renames; this had not caught
    up (raised in review 2026-08-26).

    The temporary file is a sibling so the rename stays on one filesystem,
    where it is atomic. And it is opened again before the rename: a file that
    saved without raising can still be unopenable, and finding that out after
    replacing the original is exactly too late.
    """
    from docx import Document  # python-docx

    tmp = out_path.with_name(out_path.name + f".{os.getpid()}.part")
    try:
        doc.save(str(tmp))
        Document(str(tmp))  # raises if what we just wrote is not a document
        os.replace(tmp, out_path)
    except BaseException:
        tmp.unlink(missing_ok=True)
        raise


def update_docx(doc_path: Path, out_path: Path, lines: list[str]) -> dict:
    """Apply `lines` to the document at `doc_path`, saving to `out_path`."""
    from docx import Document  # python-docx

    doc = Document(str(doc_path))
    paragraphs = doc.paragraphs
    # Blank paragraphs are layout, not content, and are never matched or moved.
    indexed = [(i, p) for i, p in enumerate(paragraphs) if p.text.strip()]
    doc_texts = [p.text for _, p in indexed]

    pairs = pair_up(doc_texts, lines)
    bounds = content_bounds(pairs)
    covered = coverage(pairs, doc_texts, lines)
    changes: list[Change] = []
    untouched = len(doc_texts)

    if bounds is None or covered < _MIN_COVERAGE:
        # Nothing recognisable. Saving an unchanged copy is the honest result;
        # rewriting on a guess would destroy the document.
        save_atomically(doc, out_path)
        return {
            "out_path": str(out_path),
            "edited": 0,
            "added": 0,
            "removed": 0,
            "untouched": untouched,
            "matched": False,
            "coverage": covered,
            "changes": [],
        }

    lo, hi = bounds
    edited = added = removed = 0
    # Deletions are applied after the walk: removing a paragraph mid-walk
    # would invalidate the indices everything else is keyed by.
    to_remove = []
    last_par = None  # where an inserted paragraph gets attached, and its style
    last_index = -1  # that paragraph's number in the document, for the report

    for pos, line_idx in pairs:
        if pos is not None and not (lo <= pos <= hi):
            continue  # cover sheet and anything past the script: not ours
        if pos is not None and line_idx is not None:
            par = indexed[pos][1]
            new_text = lines[line_idx]
            if set_paragraph_text(par, new_text):
                edited += 1
                changes.append(Change("แก้ไข", indexed[pos][0], doc_texts[pos], new_text))
            last_par, last_index = par, indexed[pos][0]
        elif pos is not None:
            par = indexed[pos][1]
            to_remove.append(par)
            removed += 1
            changes.append(Change("ตัดออก", indexed[pos][0], doc_texts[pos], ""))
        elif line_idx is not None and last_par is not None:
            # A new paragraph takes the previous one's shape, so added text
            # arrives in the document's own style rather than Word's default.
            from docx.text.paragraph import Paragraph

            new_p = copy.deepcopy(last_par._p)
            last_par._p.addnext(new_p)
            par = Paragraph(new_p, last_par._parent)
            set_paragraph_text(par, lines[line_idx])
            added += 1
            changes.append(Change("เพิ่มเติม", last_index, "", lines[line_idx]))
            last_par = par

    for par in to_remove:
        par._p.getparent().remove(par._p)

    save_atomically(doc, out_path)
    return {
        "out_path": str(out_path),
        "edited": edited,
        "added": added,
        "removed": removed,
        "untouched": untouched - edited - removed,
        "matched": True,
        "coverage": covered,
        "changes": [c.to_dict() for c in changes],
    }
